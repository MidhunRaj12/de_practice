from docx import Document
from docx.shared import Inches

# Create a new Document
doc = Document()

# Title
doc.add_heading('Senior Data Engineer Interview Preparation Guide', 0)

# Summary Section
doc.add_heading('Professional Profile Summary', level=1)
doc.add_paragraph(
    "Senior Data Engineer with 9+ years of experience specializing in cloud-native data pipelines. "
    "Expertise in Python, AWS (S3, Glue, Lambda), PySpark, SQL, dbt, Snowflake, and Airflow. "
    "Strong focus on data security (PII), performance optimization, and scalable architecture."
)

# Experience Table 1: IBM
doc.add_heading('Experience: IBM (Healthcare Migration)', level=2)
table_ibm = doc.add_table(rows=1, cols=2)
table_ibm.style = 'Table Grid'
hdr_cells = table_ibm.rows[0].cells
hdr_cells[0].text = 'Category'
hdr_cells[1].text = 'Details'

data_ibm = [
    ('Role', 'Data Engineer'),
    ('Context', 'Early-stage migration of healthcare datasets to AWS and Snowflake.'),
    ('Core Tech', 'Python, Snowflake, AWS S3, RBAC, Masking.'),
    ('Key Challenge', 'PII Leakage: Source schema changes introduced raw patient data into unmapped fields.'),
    ('The Fix', 'Implemented a Metadata-Driven Validation Layer (Quarantine Pattern).'),
    ('Result', 'Prevented HIPAA breaches by isolating non-compliant batches and triggering CloudWatch alerts.')
]

for cat, detail in data_ibm:
    row_cells = table_ibm.add_row().cells
    row_cells[0].text = cat
    row_cells[1].text = detail

# Experience Table 2: NBN Australia (Infosys)
doc.add_heading('Experience: NBN Australia (Telecom Telemetry)', level=2)
table_nbn = doc.add_table(rows=1, cols=2)
table_nbn.style = 'Table Grid'
hdr_cells = table_nbn.rows[0].cells
hdr_cells[0].text = 'Category'
hdr_cells[1].text = 'Details'

data_nbn = [
    ('Role', 'Senior Data Engineer'),
    ('Context', 'Processing massive volumes of network telemetry data.'),
    ('Core Tech', 'PySpark, AWS Glue, Airflow, S3, Snowflake.'),
    ('Key Challenge', 'Small File Problem & Data Skew: Thousands of 10KB files slowed ingestion; skewed partitions caused OOM.'),
    ('The Fix', 'Salting for skew; Compaction DAG using coalesce() for small files.'),
    ('Result', 'Query performance improved by 3x; ingestion cost reduced by 20%.')
]

for cat, detail in data_nbn:
    row_cells = table_nbn.add_row().cells
    row_cells[0].text = cat
    row_cells[1].text = detail

# Experience Table 3: Infosys
doc.add_heading('Experience: Infosys (Data & Automation)', level=2)
table_infosys = doc.add_table(rows=1, cols=2)
table_infosys.style = 'Table Grid'
hdr_cells = table_infosys.rows[0].cells
hdr_cells[0].text = 'Category'
hdr_cells[1].text = 'Details'

data_infosys = [
    ('Role', 'Team Lead - Data & Automation'),
    ('Context', 'ETL utility development and legacy system migrations.'),
    ('Core Tech', 'Python, PostgreSQL, MySQL, AWS RDS.'),
    ('Key Challenge', 'Performance Bottlenecks: Critical reporting queries taking 40+ minutes due to implicit casting.'),
    ('The Fix', 'SQL Refactoring (CTEs, Late Row Lookups) and Data Quality Gatekeepers.'),
    ('Result', 'Query runtime reduced from 40 mins to <2 mins; improved data consistency.')
]

for cat, detail in data_infosys:
    row_cells = table_infosys.add_row().cells
    row_cells[0].text = cat
    row_cells[1].text = detail

# Top Interview Questions Section
doc.add_page_break()
doc.add_heading('Detailed Interview Questions & Strategic Answers', level=1)

# Question 1
doc.add_heading('1. Role and Architecture (The "Gatekeeper" Pipeline)', level=2)
doc.add_paragraph(
    "Question: Can you walk us through the architecture of your most recent production pipeline?\n"
    "Strategic Answer: Describe the E2E flow. Mention S3 for landing, Airflow for orchestration, "
    "and Glue for the 'Gatekeeper' logic. Emphasize why you chose a modular approach: 'By separating "
    "the PII scan from the load process, we ensured that only validated, compliant data reached Snowflake.'"
)

# Question 2
doc.add_heading('2. Handling Technical Failure (Data Skew)', level=2)
doc.add_paragraph(
    "Question: How do you handle a scenario where a Spark job is failing with OOM due to skewed data?\n"
    "Strategic Answer: Explain Salting. 'I identify the skewed key and add a random integer to it "
    "to force Spark to redistribute the data across executors.' Also mention modern Spark features like "
    "Adaptive Query Execution (AQE) with skewJoin.enabled=true."
)

# Question 3
doc.add_heading('3. Data Security & Compliance', level=2)
doc.add_paragraph(
    "Question: How do you manage PII in a multi-tenant or multi-role environment like Snowflake?\n"
    "Strategic Answer: Talk about the Quarantine Pattern for the ETL layer and Snowflake RBAC for "
    "the warehouse layer. 'I use Dynamic Data Masking policies so that unauthorized roles see "
    "masked strings (e.g., ###-##-XXXX) while authorized roles see the raw data.'"
)

# Question 4
doc.add_heading('4. Performance Optimization (The "Small File" Issue)', level=2)
doc.add_paragraph(
    "Question: What is the 'Small File Problem' in S3 and how do you resolve it?\n"
    "Strategic Answer: Explain that too many files create metadata overhead. 'My solution was a "
    "compaction job in PySpark that reads the fragmented data and uses coalesce() to write it back "
    "in 128MB-256MB Parquet blocks, which are the optimal size for S3 and Snowflake reads.'"
)

# Question 5
doc.add_heading('5. SQL & Data Quality', level=2)
doc.add_paragraph(
    "Question: How do you ensure your transformations are both correct and performant?\n"
    "Strategic Answer: Discuss dbt-tests and Row Count Variance. 'I implement a variance check "
    "in Airflow. If a batch deviates by >20% from the rolling average, the pipeline fails-fast. "
    "For performance, I avoid implicit casting by ensuring join keys have matching data types.'"
)

# Save the document
file_path = r"C:\Users\midhu\OneDrive\Documents\Interview Q SS\Midhun_Raj_Interview_Prep.docx"
doc.save(file_path)